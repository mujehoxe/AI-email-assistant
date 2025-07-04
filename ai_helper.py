import streamlit as st
import json
import subprocess
import time
import threading
import os
import traceback
import re
import requests
import openai
from typing import Dict, List, Any, Optional, Union, Tuple
import io

# Try importing document libraries
try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Try importing PDF libraries
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import pdfkit
    PDFKIT_AVAILABLE = True
except ImportError:
    PDFKIT_AVAILABLE = False

# Configuration
G4F_API_HOST = os.environ.get("G4F_API_HOST", "localhost")
G4F_API_PORT = int(os.environ.get("G4F_API_PORT", "8080"))
G4F_API_BASE_URL = f"http://{G4F_API_HOST}:{G4F_API_PORT}/v1"

# Default model and fallback models
DEFAULT_MODEL = "llama-4-scout"
FALLBACK_MODELS = [
    "llama-4-scout",  # Default
    "gpt-4o",
    "gpt-3.5-turbo",
    "gpt-4",
    "claude-3-opus",
    "claude-3-sonnet",
    "gemini-pro",
    "mistral-medium",
]

# Global state
api_server_running = False
request_in_progress = False
cancel_request = False
AVAILABLE_MODELS = FALLBACK_MODELS.copy()

# Type definitions
ContactInfo = Dict[str, str]
TemplateData = Dict[str, str]
APIResponse = Dict[str, Any]
ResumeData = Dict[str, str]


def extract_email_from_text(text: str) -> str:
    """Extract email addresses from text"""
    email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    emails = re.findall(email_pattern, text)
    return emails[0] if emails else ""


def extract_company_name(job_description: str) -> str:
    """Extract company name from job description"""
    patterns = [
        r"(?:at|with|for|join)\s+([\w\s&\-\.]+?)(?:is\s+looking|is\s+seeking|is\s+hiring|is\s+searching|has\s+an\s+opening|has\s+a\s+job|has\s+a\s+position|has\s+an\s+opportunity)",
        r"(?:at|with|for|join)\s+([\w\s&\-\.]+?)(?:\.|,|\sin\s)",
        r"([\w\s&\-\.]+?)(?:\sis\s+looking|\sis\s+seeking|\sis\s+hiring|\shas\s+an\s+opening|\shas\s+a\s+job|\shas\s+a\s+position)",
        r"(?:company|employer):\s*([\w\s&\-\.]+)",
        r"(?:about\s+us|about\s+the\s+company|company\s+overview|about\s+the\s+team)\s*(?:\n|\r\n?)([\w\s&\-\.]+)",
        r"(?:about\s+)([\w\s&\-\.]+)(?:\s+[\w\s&\-\.]+\s+is\s+a)",
    ]

    for pattern in patterns:
        matches = re.search(pattern, job_description, re.IGNORECASE)
        if matches:
            company_name = matches.group(1).strip()
            company_name = re.sub(r"\s+", " ", company_name)
            if company_name:
                return company_name

    # If no company name found, try to extract from emails in the job description
    emails = extract_email_from_text(job_description)
    if emails:
        # Try to get company name from email domain
        email = emails if isinstance(emails, str) else emails[0]
        domain = email.split('@')[-1]
        if domain:
            # Remove TLD (.com, .org, etc.) and convert to title case
            company = domain.split('.')[0].title()
            if company and len(company) > 2:  # Ensure it's a meaningful name
                return company

    return ""


def extract_contact_info(job_description: str) -> ContactInfo:
    """Extract contact information from job description

    This function only extracts information from the job description, not the resume.
    It returns a dictionary with email, name, and company extracted from the job description.
    """
    email = extract_email_from_text(job_description)
    company_name = extract_company_name(job_description)
    name = ""

    # Extract name based on email if available
    name_patterns = []
    if email:
        name_patterns.extend(
            [
                r"(?:contact|send|email|apply|resume to|cv to|application to)[:\s]+([\w\s]+)[\s\(<]+"
                + re.escape(email),
                r"([\w\s]+)[\s\(<]+" + re.escape(email),
            ]
        )

    # General name patterns
    name_patterns.extend(
        [
            r"(?:contact|send|email|apply|resume to|cv to|application to)[:\s]+([\w\s]+)",
            r"(?:recruiter|hiring manager|hr manager|contact person)[:\s]+([\w\s]+)",
        ]
    )

    # Try to find name in job description
    for pattern in name_patterns:
        matches = re.search(pattern, job_description, re.IGNORECASE)
        if matches:
            name = matches.group(1).strip()
            break

    # If name not found in job description but we have email, extract from email
    if not name and email:
        # Try to extract name from the email address
        local_part = email.split('@')[0]
        
        # Try different common email formats
        if '.' in local_part:  # firstname.lastname@domain.com
            parts = local_part.split('.')
            name = ' '.join(part.title() for part in parts)
        elif '_' in local_part:  # firstname_lastname@domain.com
            parts = local_part.split('_')
            name = ' '.join(part.title() for part in parts)
        elif '-' in local_part:  # firstname-lastname@domain.com
            parts = local_part.split('-')
            name = ' '.join(part.title() for part in parts)
        else:
            # Try to detect if it's a name by checking if it contains digits
            if not any(c.isdigit() for c in local_part):
                name = local_part.title()
        
        # If we couldn't extract a name from email, fall back to utils.extract_name
        if not name:
            from utils import extract_name
            name = extract_name(email)

    # Filter out generic or invalid names
    generic_names = [
        "at",
        "hr",
        "info",
        "jobs",
        "careers",
        "admin",
        "recruitment",
        "recruiting",
        "contact",
        "apply",
        "job",
        "resume",
        "application",
        "hiring",
        "talent",
        "human resources",
        "humanresources",
    ]

    # Check if the extracted name is generic
    if name.lower() in generic_names or len(name) <= 2:
        name = ""  # Clear the name if it's generic

    # If company name is not found but we have email, try to extract from email domain
    if not company_name and email:
        domain = email.split('@')[-1]
        if domain:
            # Remove TLD (.com, .org, etc.) and convert to title case
            company = domain.split('.')[0].title()
            if company and len(company) > 2:  # Ensure it's a meaningful name
                company_name = company

    return {"email": email, "name": name, "company": company_name}


def check_api_server_status() -> bool:
    """Check if the G4F API server is running"""
    try:
        response = requests.get(
            f"http://{G4F_API_HOST}:{G4F_API_PORT}/v1/models", timeout=5
        )
        return response.status_code == 200
    except (requests.RequestException, ConnectionError):
        return False


def get_available_models() -> List[str]:
    """Get list of available models from G4F API"""
    if not check_api_server_status():
        return FALLBACK_MODELS.copy()

    try:
        response = requests.get(
            f"http://{G4F_API_HOST}:{G4F_API_PORT}/v1/models", timeout=5
        )

        if response.status_code == 200:
            data = response.json()

            if "data" in data and isinstance(data["data"], list):
                models = []
                for model in data["data"]:
                    if isinstance(model, dict) and "id" in model:
                        model_id = model["id"]
                        if model_id != "default" and model_id not in models:
                            models.append(model_id)

                if models:
                    return models

        return FALLBACK_MODELS.copy()

    except Exception as e:
        st.warning(f"Could not fetch available models from G4F API: {str(e)}")
        return FALLBACK_MODELS.copy()


def start_g4f_api_server() -> None:
    """Start the G4F API server in a separate thread"""
    global api_server_running, AVAILABLE_MODELS

    if api_server_running:
        return

    def run_server():
        try:
            st.info("Starting G4F API server...")
            process = subprocess.Popen(
                ["python", "-c", "from g4f.api import run_api; run_api()"],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )

            for line in process.stderr:
                st.error(f"G4F API server error: {line.decode().strip()}")
            for line in process.stdout:
                st.info(f"G4F API server: {line.decode().strip()}")

        except Exception as e:
            st.error(f"Failed to start G4F API server: {str(e)}")
            st.error(traceback.format_exc())

    # Start server in a separate thread
    server_thread = threading.Thread(target=run_server)
    server_thread.daemon = True
    server_thread.start()

    # Wait for server to start
    time.sleep(5)
    api_server_running = True

    # Update session state
    st.session_state.api_server_status = "Running"

    # Update available models
    time.sleep(2)  # Wait for server initialization
    try:
        fresh_models = get_available_models()
        if fresh_models:
            AVAILABLE_MODELS = fresh_models
            if "ai_models" in st.session_state:
                st.session_state.ai_models = fresh_models
    except Exception as e:
        st.warning(f"Could not update model list: {e}")


def get_ai_client():
    """Get OpenAI-compatible client configured to use G4F API"""
    base_url = f"http://{G4F_API_HOST}:{G4F_API_PORT}/v1"

    try:
        client = openai.OpenAI(
            api_key="not-needed",
            base_url=base_url,
        )
        return client
    except Exception as e:
        try:
            import g4f

            return g4f
        except ImportError:
            st.error("Failed to import g4f. Make sure g4f is installed.")
            return None


def get_g4f_model_from_name(model_name: str):
    """Map OpenAI model name to g4f provider if possible"""
    try:
        import g4f
        from g4f.models import (
            gpt_4,
            gpt_4_turbo,
            claude_3_opus,
            claude_3_sonnet,
            gemini_pro,
        )
        
        model_map = {
            "gpt-4o": gpt_4_turbo,  # Using gpt_4_turbo as fallback for gpt-4o
            "gpt-3.5-turbo": gpt_4,  # Using gpt_4 as fallback since gpt_35_turbo doesn't exist
            "gpt-4": gpt_4,
            "gpt-4-turbo": gpt_4_turbo,
            "claude-3-opus": claude_3_opus,
            "claude-3-sonnet": claude_3_sonnet,
            "gemini-pro": gemini_pro,
        }

        return model_map.get(model_name, gpt_4)  # Default to gpt_4 if model not found
    except ImportError:
        st.error("Failed to import g4f. Make sure g4f is installed.")
        return None


def create_prompt(
    job_description: str,
    resume_content: str,
    current_template: Dict[str, str],
    contact_info: ContactInfo,
) -> str:
    """Create prompt for AI model"""
    # Extract first name if full name is available
    first_name = ""
    last_name = ""
    if contact_info["name"]:
        name_parts = contact_info["name"].split()
        if name_parts:
            first_name = name_parts[0]
            if len(name_parts) > 1:
                last_name = name_parts[-1]

    return f"""
    Create a personalized email template for a job application based on this information:
    
    ## Job Description:
    {job_description}
    
    ## My Resume:
    {resume_content}
    
    ## Current Email Template:
    Greeting: {current_template.get('greeting', '')}
    Body: {current_template.get('body', '')}
    Signature: {current_template.get('signature', '')}
    
    ## Contact Information (extracted from job description):
    Email: {contact_info['email']}
    Name: {contact_info['name']}
    First Name: {first_name}
    Last Name: {last_name}
    Company: {contact_info['company'] or 'Not found (which is okay)'}
    
    Please create:
    1. A brief, conversational email that sounds like a real person wrote it
    2. Highlight 2-3 most relevant skills/experiences matching the job requirements
    3. Keep it concise and enthusiastic
    4. Maintain the greeting/body/signature structure
    5. In the greeting:
       - If a valid recipient name is found, use their first name (preferred) or last name with appropriate title (Mr./Ms.)
       - If no valid name is found (or name is generic like "at", "hr", etc.), use "Hiring Manager" in the greeting
       - DO NOT use a full name in the greeting
    6. Keep the placeholder {{{{name}}}} in the greeting, but make sure the instructions specify to use first or last name only
    7. Keep the placeholder {{{{position}}}} where appropriate
    8. If a company name was detected, mention it in the body. If no company name was found, try to extract it from the email domain (e.g., "acme" from "jobs@acme.com")
    
    Important: 
    - Only extract information from the job description, not the resume
    - Use the resume only to match skills and experience to the job requirements
    - If you don't have a company name, try to extract it from the email domain if available
    - If you can't determine a company name, leave the company field empty. DO NOT use placeholder text like "the company" or "[Company Name]"
    - If the recipient name is missing, try to extract it from the email address (e.g., "John Smith" from "john.smith@company.com")
    - If the recipient name is still missing, invalid, or generic (like "at", "hr", "info", "jobs", "careers", "admin", etc.), leave the recipient_name field empty in your response and use "Hiring Manager" in the greeting
    
    Format your response as a valid JSON object with the following structure:
    
    {{
      "greeting": "Your greeting here",
      "body": "Your email body here",
      "signature": "Your signature here",
      "position": "Extracted position from job description",
      "employer": "Extracted employer name from job description or email domain if found",
      "subject": "Suggested email subject",
      "recipient_name": "Extracted recipient name from job description or email if found (leave empty if invalid or generic)",
      "recipient_email": "Extracted email from job description if found"
    }}
    """


def parse_ai_response(
    result: str, contact_info: ContactInfo
) -> Tuple[bool, Dict[str, Any]]:
    """Parse the AI response and extract template data

    Uses contact_info from the job description as fallback if the AI doesn't provide values.
    """
    # Try to parse the JSON directly
    try:
        template = json.loads(result)

        # Check if template has all required fields
        required_fields = ["greeting", "body", "signature"]
        if all(field in template for field in required_fields):
            return True, {"success": True, "template": template}
        else:
            missing = [field for field in required_fields if field not in template]
            return False, {
                "success": False,
                "error": f"Response missing required fields: {', '.join(missing)}",
                "template": template,
            }

    except json.JSONDecodeError:
        # Try to extract JSON from markdown
        try:
            json_pattern = r"```(?:json)?\s*([\s\S]*?)\s*```"
            match = re.search(json_pattern, result)

            if match:
                json_str = match.group(1)
                template = json.loads(json_str)

                required_fields = ["greeting", "body", "signature"]
                if all(field in template for field in required_fields):
                    return True, {"success": True, "template": template}
                else:
                    missing = [
                        field for field in required_fields if field not in template
                    ]
                    return False, {
                        "success": False,
                        "error": f"Response missing required fields: {', '.join(missing)}",
                        "template": template,
                    }
        except:
            pass

        # Manual extraction as last resort
        st.warning("JSON parsing failed. Trying manual extraction...")

        lines = result.split("\n")
        template = {
            "greeting": "",
            "body": "",
            "signature": "",
            "position": "",
            "employer": "",
            "subject": "",
            "recipient_email": contact_info["email"],
            "recipient_name": contact_info["name"],
        }

        current_section = None

        for line in lines:
            line = line.strip()
            if line.lower().startswith("greeting:"):
                current_section = "greeting"
                template["greeting"] = line[len("greeting:") :].strip()
            elif line.lower().startswith("body:"):
                current_section = "body"
                template["body"] = line[len("body:") :].strip()
            elif line.lower().startswith("signature:"):
                current_section = "signature"
                template["signature"] = line[len("signature:") :].strip()
            elif line.lower().startswith("position:"):
                template["position"] = line[len("position:") :].strip()
            elif line.lower().startswith("employer:"):
                template["employer"] = line[len("employer:") :].strip()
            elif line.lower().startswith("company:"):
                if not template["employer"]:
                    template["employer"] = line[len("company:") :].strip()
            elif line.lower().startswith("subject:"):
                template["subject"] = line[len("subject:") :].strip()
            elif line.lower().startswith("recipient_email:"):
                template["recipient_email"] = line[len("recipient_email:") :].strip()
            elif line.lower().startswith("recipient_name:"):
                template["recipient_name"] = line[len("recipient_name:") :].strip()
            elif current_section == "greeting":
                template["greeting"] += "\n" + line
            elif current_section == "body":
                template["body"] += "\n" + line
            elif current_section == "signature":
                template["signature"] += "\n" + line

        # Check if we have the required fields
        if template["greeting"] and template["body"] and template["signature"]:
            return True, {"success": True, "template": template}
        else:
            missing = []
            if not template["greeting"]:
                missing.append("greeting")
            if not template["body"]:
                missing.append("body")
            if not template["signature"]:
                missing.append("signature")

            return False, {
                "success": False,
                "error": f"Failed to extract required fields: {', '.join(missing)}",
                "template": template,
            }


def generate_improved_template(
    job_description: str,
    resume_content: str,
    current_template: Dict[str, str],
    model: str = DEFAULT_MODEL,
) -> Dict[str, Any]:
    """Generate an improved template based on job description and resume"""
    global request_in_progress, cancel_request

    try:
        # Set request in progress flag
        request_in_progress = True
        cancel_request = False

        # Extract contact info from job description
        contact_info = extract_contact_info(job_description)

        # Check API server status
        api_server_running = check_api_server_status()

        # Create prompt
        prompt = create_prompt(
            job_description, resume_content, current_template, contact_info
        )

        # Get client for API request
        client = get_ai_client()

        # Generate response
        result = ""
        if api_server_running and hasattr(client, "chat"):
            # Using OpenAI-compatible API
            st.info(f"Sending request to AI model: {model}...")

            try:
                response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant."},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.7,
                )

                result = response.choices[0].message.content
            except Exception as api_error:
                st.error(f"API error: {str(api_error)}")
                # Fall back to direct g4f
                api_server_running = False

        if not api_server_running:
            # Using g4f directly as fallback
            st.warning("Using g4f directly (this may take longer)...")

            try:
                import g4f

                g4f_model = get_g4f_model_from_name(model)
                if not g4f_model:
                    return {"success": False, "error": "Failed to get g4f model"}

                result = g4f.ChatCompletion.create(
                    model=g4f_model,
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant."},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.7,
                )
            except Exception as g4f_error:
                st.error(f"Error using g4f directly: {str(g4f_error)}")
                return {"success": False, "error": str(g4f_error)}

        # Check if request was cancelled
        if cancel_request:
            return {"success": False, "error": "Request cancelled by user"}

        # Parse the response
        success, response_data = parse_ai_response(result, contact_info)
        return response_data

    except Exception as e:
        st.error(f"Error in generate_improved_template: {str(e)}")
        st.error(traceback.format_exc())
        return {"success": False, "error": str(e)}
    finally:
        # Reset request in progress flag
        request_in_progress = False


def modify_original_word_document(
    original_file_bytes: bytes, 
    customization_text: str, 
    original_filename: str = ""
) -> Optional[io.BytesIO]:
    """Modify the original Word document by replacing its content while preserving the original formatting"""
    if not DOCX_AVAILABLE:
        st.warning("python-docx not available. Cannot modify Word document.")
        return None
    
    try:
        # Load the original document to preserve all formatting
        doc = Document(io.BytesIO(original_file_bytes))
        
        # Get the original text to compare
        original_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Split customized text into lines
        customized_lines = [line.strip() for line in customization_text.strip().split('\n') if line.strip()]
        
        # Try to intelligently replace content while preserving formatting
        # This approach modifies existing paragraphs rather than clearing everything
        
        # Convert customized text to paragraph-by-paragraph mapping
        customized_paragraphs = []
        current_para = ""
        
        for line in customized_lines:
            if line.strip():
                if current_para and (line.endswith(':') or any(keyword in line.lower() for keyword in ['experience', 'education', 'skills', 'summary', 'professional'])):
                    # This looks like a section header, finish current paragraph
                    if current_para.strip():
                        customized_paragraphs.append(current_para.strip())
                    current_para = line
                else:
                    if current_para:
                        current_para += " " + line
                    else:
                        current_para = line
        
        # Add the last paragraph
        if current_para.strip():
            customized_paragraphs.append(current_para.strip())
        
        # If we have fewer paragraphs in the customized version, extend to match
        original_paragraph_count = len(doc.paragraphs)
        customized_paragraph_count = len(customized_paragraphs)
        
        # Method 1: Try to replace paragraph by paragraph, preserving formatting
        if customized_paragraph_count > 0:
            for i, paragraph in enumerate(doc.paragraphs[:]):
                if i < len(customized_paragraphs):
                    # Replace the text while keeping the formatting
                    new_text = customized_paragraphs[i]
                    
                    # Clear the paragraph text but keep the formatting
                    if paragraph.runs:
                        # Keep the first run's formatting and replace its text
                        first_run = paragraph.runs[0]
                        first_run.text = new_text
                        
                        # Remove any additional runs
                        for run in paragraph.runs[1:]:
                            run.text = ""
                    else:
                        # If no runs exist, add text with default formatting
                        paragraph.add_run(new_text)
                else:
                    # Clear paragraphs that don't have corresponding content
                    for run in paragraph.runs[:]:
                        run.text = ""
            
            # If we have more customized content than original paragraphs, add new paragraphs
            if customized_paragraph_count > original_paragraph_count:
                for i in range(original_paragraph_count, customized_paragraph_count):
                    new_para = doc.add_paragraph(customized_paragraphs[i])
        
        # Method 2: If the above doesn't work well, fall back to simple replacement
        else:
            # Fallback: Replace all content with customized text
            # Clear existing content
            for paragraph in doc.paragraphs[:]:
                for run in paragraph.runs[:]:
                    run.text = ""
            
            # Add customized content
            if doc.paragraphs:
                # Use the first paragraph for the customized content
                doc.paragraphs[0].add_run(customization_text)
            else:
                # Add a new paragraph if none exist
                doc.add_paragraph(customization_text)
        
        # Save to BytesIO
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer
        
    except Exception as e:
        st.error(f"Error modifying Word document: {str(e)}")
        st.error(traceback.format_exc())
        return None


def get_resume_file_extension(filename: str) -> str:
    """Get file extension from filename"""
    if not filename:
        return ""
    return filename.split('.')[-1].lower() if '.' in filename else ""


def cancel_ai_request() -> bool:
    """Cancel an in-progress AI request"""
    global cancel_request, request_in_progress

    if request_in_progress:
        cancel_request = True
        return True
    return False


def is_request_in_progress() -> bool:
    """Check if a request is in progress"""
    global request_in_progress
    return request_in_progress


def create_resume_modification_prompt(
    original_resume: str,
    job_description: str,
) -> str:
    """Create prompt for AI model to suggest specific modifications to resume"""
    return f"""
    Based on the original resume and job description provided, suggest specific modifications to better align the resume with the job requirements while maintaining truthfulness and accuracy.

    ## Original Resume:
    {original_resume}
    
    ## Job Description:
    {job_description}
    
    Please analyze the resume and suggest specific modifications in the following format:
    
    **IMPORTANT**: Do NOT generate a complete new resume. Instead, provide specific, actionable modifications that can be applied to the original text.
    
    For each suggested change, specify:
    1. **Action**: ADD, REPLACE, REMOVE, or REORDER
    2. **Location**: The specific text or section where the change should be applied
    3. **Original**: The exact text to be modified (if REPLACE or REMOVE)
    4. **New**: The replacement or additional text (if ADD or REPLACE)
    5. **Reason**: Why this change improves alignment with the job requirements
    
    **PRIORITY MODIFICATIONS (Must Include):**
    1. **Job Title Alignment**: If the resume doesn't have a job title that matches the position, ADD or REPLACE the current title with one that closely aligns with the job posting (e.g., "Software Engineer" → "Full Stack Developer" if applying for a Full Stack position)
    2. **Professional Summary**: ADD a compelling 2-3 sentence professional summary at the top that highlights the candidate's most relevant experience and skills for this specific role. This should be tailored to the job description keywords and requirements.
    
    **Additional Guidelines:**
    - Only suggest changes that enhance existing information from the original resume
    - Do not fabricate experience, skills, or achievements
    - Focus on keyword optimization, emphasis, and reorganization
    - Suggest adding relevant skills that are already implied by the candidate's experience
    - Recommend reordering sections or bullet points to highlight most relevant experience first
    - Ensure the job title reflects the candidate's actual capabilities based on their experience
    - Make the professional summary compelling and specific to the role, highlighting quantifiable achievements when available
    
    Format your response as a JSON object with the following structure:
    
    {{
      "modifications": [
        {{
          "action": "REPLACE|ADD|REMOVE|REORDER",
          "location": "Section name or specific text to locate the change",
          "original": "Exact text to be replaced or removed (if applicable)",
          "new": "New text to add or replace with (if applicable)",
          "reason": "Explanation of why this change improves job alignment"
        }}
      ],
      "summary": "Brief overview of the suggested changes and their impact"
    }}
    
    Example modification:
    {{
      "action": "REPLACE",
      "location": "Professional Summary",
      "original": "Software Engineer with 5 years of experience",
      "new": "Full Stack Developer with 5 years of experience in React, Node.js, and cloud technologies",
      "reason": "Uses job title from posting and highlights specific technologies mentioned in requirements"
    }}
    """


def create_pdf_from_text_reportlab(text_content: str, original_filename: str = "") -> Optional[io.BytesIO]:
    """Create a professional PDF from text using ReportLab"""
    if not REPORTLAB_AVAILABLE:
        st.warning("ReportLab not available. Cannot create PDF document.")
        return None
    
    try:
        # Create buffer
        buffer = io.BytesIO()
        
        # Create document
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor='black'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=12,
            textColor='black'
        )
        
        # Parse content
        lines = [line.strip() for line in text_content.strip().split('\n') if line.strip()]
        
        # Extract name and job title
        name = ""
        job_title = ""
        content_start_idx = 0
        
        for i, line in enumerate(lines[:3]):
            if i == 0:
                if any(keyword in line.lower() for keyword in ['engineer', 'developer', 'manager', 'analyst', 'specialist']):
                    job_title = line
                    content_start_idx = 1
                else:
                    name = line
                    content_start_idx = 1
            elif i == 1 and not job_title:
                if any(keyword in line.lower() for keyword in ['engineer', 'developer', 'manager', 'analyst', 'specialist']):
                    job_title = line
                    content_start_idx = 2
                else:
                    content_start_idx = 1
                    break
        
        # Add name as title
        if name:
            story.append(Paragraph(name, title_style))
        
        # Add job title
        if job_title:
            story.append(Paragraph(job_title, styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Process remaining content
        for line in lines[content_start_idx:]:
            line = line.strip()
            if not line:
                continue
            
            # Check if section header
            if (line.endswith(':') and len(line) < 60 and 
                any(keyword in line.lower() for keyword in ['experience', 'education', 'skills', 'summary'])):
                story.append(Spacer(1, 12))
                story.append(Paragraph(line, heading_style))
            else:
                story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        st.error(f"Error creating PDF with ReportLab: {str(e)}")
        return None


def modify_pdf_with_pymupdf(
    original_file_bytes: bytes, 
    customization_text: str, 
    original_filename: str = ""
) -> Optional[io.BytesIO]:
    """Modify PDF using PyMuPDF by replacing content while preserving some layout"""
    if not PYMUPDF_AVAILABLE:
        st.warning("PyMuPDF not available. Cannot modify PDF document.")
        return create_pdf_from_text_reportlab(customization_text, original_filename)
    
    try:
        # Open original PDF to get styling context
        original_doc = fitz.open(stream=original_file_bytes, filetype="pdf")
        
        # Create new PDF document
        new_doc = fitz.open()  # Create empty PDF
        
        # Get page size from original (if available)
        page_width = 612  # Default letter width
        page_height = 792  # Default letter height
        
        if len(original_doc) > 0:
            original_page = original_doc[0]
            page_rect = original_page.rect
            page_width = page_rect.width
            page_height = page_rect.height
        
        # Create new page
        new_page = new_doc.new_page(width=page_width, height=page_height)
        
        # Parse customized text
        lines = [line.strip() for line in customization_text.strip().split('\n') if line.strip()]
        
        # Starting position
        y_position = page_height - 72  # Start 1 inch from top
        margin = 72  # 1 inch margin
        line_height = 14
        
        # Extract name and job title
        name = ""
        job_title = ""
        content_start_idx = 0
        
        for i, line in enumerate(lines[:3]):
            if i == 0:
                if any(keyword in line.lower() for keyword in ['engineer', 'developer', 'manager', 'analyst', 'specialist']):
                    job_title = line
                    content_start_idx = 1
                else:
                    name = line
                    content_start_idx = 1
            elif i == 1 and not job_title:
                if any(keyword in line.lower() for keyword in ['engineer', 'developer', 'manager', 'analyst', 'specialist']):
                    job_title = line
                    content_start_idx = 2
                else:
                    content_start_idx = 1
                    break
        
        # Add name (larger, centered)
        if name:
            text_rect = fitz.Rect(margin, y_position - 20, page_width - margin, y_position)
            new_page.insert_text(
                (page_width/2, y_position), 
                name, 
                fontsize=18, 
                fontname="helv-bold",
                color=(0, 0, 0)
            )
            y_position -= 30
        
        # Add job title (centered, smaller)
        if job_title:
            new_page.insert_text(
                (page_width/2, y_position), 
                job_title, 
                fontsize=14, 
                fontname="helv-oblique",
                color=(0, 0, 0)
            )
            y_position -= 25
        
        # Add remaining content
        for line in lines[content_start_idx:]:
            line = line.strip()
            if not line:
                continue
            
            # Check if we need a new page
            if y_position < 72:  # 1 inch from bottom
                new_page = new_doc.new_page(width=page_width, height=page_height)
                y_position = page_height - 72
            
            # Check if section header
            if (line.endswith(':') and len(line) < 60 and 
                any(keyword in line.lower() for keyword in ['experience', 'education', 'skills', 'summary'])):
                y_position -= 10  # Extra space before section
                new_page.insert_text(
                    (margin, y_position), 
                    line, 
                    fontsize=12, 
                    fontname="helv-bold",
                    color=(0, 0, 0)
                )
                y_position -= line_height
            else:
                # Regular text
                new_page.insert_text(
                    (margin, y_position), 
                    line, 
                    fontsize=10, 
                    fontname="helv",
                    color=(0, 0, 0)
                )
                y_position -= line_height
        
        # Save to buffer
        buffer = io.BytesIO()
        new_doc.save(buffer)
        new_doc.close()
        original_doc.close()
        
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        st.error(f"Error modifying PDF with PyMuPDF: {str(e)}")
        # Fall back to ReportLab
        return create_pdf_from_text_reportlab(customization_text, original_filename)


def modify_original_resume_document(
    original_file_bytes: bytes, 
    customization_text: str, 
    original_filename: str = ""
) -> Optional[io.BytesIO]:
    """Modify the original resume document while preserving styling"""
    file_extension = get_resume_file_extension(original_filename)
    
    if file_extension == "docx":
        return modify_original_word_document(original_file_bytes, customization_text, original_filename)
    elif file_extension == "pdf":
        # Try PyMuPDF first, fall back to ReportLab if needed
        return modify_pdf_with_pymupdf(original_file_bytes, customization_text, original_filename)
    else:
        st.warning(f"Unsupported file format: {file_extension}. Only PDF and DOCX are supported for customization.")
        return None


def parse_resume_modifications(ai_response: str) -> Dict[str, Any]:
    """Parse the AI response containing resume modifications"""
    try:
        # Try to parse the JSON directly
        modifications_data = json.loads(ai_response)
        
        # Validate the structure
        if "modifications" in modifications_data and isinstance(modifications_data["modifications"], list):
            return {"success": True, "data": modifications_data}
        else:
            return {"success": False, "error": "Invalid modifications format"}
            
    except json.JSONDecodeError:
        # Try to extract JSON from markdown
        try:
            json_pattern = r"```(?:json)?\s*([\s\S]*?)\s*```"
            match = re.search(json_pattern, ai_response)
            
            if match:
                json_str = match.group(1)
                # Remove trailing commas before parsing JSON
                json_str = re.sub(r",\s*([}\]])", r"\1", json_str)
                # Replace null-like entries for "original" and "new" with empty strings  
                json_str = re.sub(r'"original"\s*:\s*,', '"original": "",', json_str)
                json_str = re.sub(r'"new"\s*:\s*,', '"new": "",', json_str)
                
                modifications_data = json.loads(json_str)
                
                if "modifications" in modifications_data and isinstance(modifications_data["modifications"], list):
                    return {"success": True, "data": modifications_data}
                else:
                    return {"success": False, "error": "Invalid modifications format in extracted JSON"}
            else:
                return {"success": False, "error": "No JSON found in response"}
                
        except json.JSONDecodeError:
            return {"success": False, "error": "Failed to parse JSON from response"}


def apply_resume_modifications_to_document(
    original_file_bytes: bytes, 
    original_text: str, 
    modifications: List[Dict[str, str]], 
    filename: str
) -> Tuple[Optional[io.BytesIO], List[str]]:
    """Apply modifications directly to the original document while preserving formatting"""
    applied_changes = []
    
    try:
        file_extension = get_resume_file_extension(filename)
        
        if file_extension == "pdf":
            return apply_modifications_to_pdf(original_file_bytes, original_text, modifications, filename)
        elif file_extension == "docx":
            return apply_modifications_to_docx(original_file_bytes, original_text, modifications, filename)
        else:
            applied_changes.append(f"❌ Unsupported file format: {file_extension}")
            return None, applied_changes
            
    except Exception as e:
        applied_changes.append(f"❌ ERROR applying modifications to document: {str(e)}")
        return None, applied_changes


def apply_modifications_to_docx(
    original_file_bytes: bytes, 
    original_text: str, 
    modifications: List[Dict[str, str]], 
    filename: str
) -> Tuple[Optional[io.BytesIO], List[str]]:
    """Apply modifications to DOCX document while preserving formatting"""
    if not DOCX_AVAILABLE:
        return None, ["❌ python-docx not available"]
    
    applied_changes = []
    
    try:
        # Load the original document
        doc = Document(io.BytesIO(original_file_bytes))
        
        # Sort modifications by action priority
        action_priority = {"REMOVE": 1, "REPLACE": 2, "ADD": 3, "REORDER": 4}
        sorted_modifications = sorted(modifications, key=lambda x: action_priority.get(x.get("action", ""), 5))
        
        # Apply modifications
        for mod in sorted_modifications:
            action = mod.get("action", "").upper()
            original_text_to_find = mod.get("original", "")
            new_text = mod.get("new", "")
            reason = mod.get("reason", "")
            
            if action == "REPLACE" and original_text_to_find and new_text:
                modified = False
                # Search through all paragraphs
                for paragraph in doc.paragraphs:
                    if original_text_to_find in paragraph.text:
                        # Preserve formatting by replacing text while keeping runs
                        full_text = paragraph.text
                        new_full_text = full_text.replace(original_text_to_find, new_text)
                        
                        # Clear paragraph and add new text
                        paragraph.clear()
                        paragraph.add_run(new_full_text)
                        
                        applied_changes.append(f"✅ REPLACED in document: '{original_text_to_find[:30]}...' with '{new_text[:30]}...'")
                        modified = True
                        break
                
                # Also check tables
                if not modified:
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if original_text_to_find in cell.text:
                                    for paragraph in cell.paragraphs:
                                        if original_text_to_find in paragraph.text:
                                            full_text = paragraph.text
                                            new_full_text = full_text.replace(original_text_to_find, new_text)
                                            paragraph.clear()
                                            paragraph.add_run(new_full_text)
                                            applied_changes.append(f"✅ REPLACED in table: '{original_text_to_find[:30]}...' with '{new_text[:30]}...'")
                                            modified = True
                                            break
                                if modified:
                                    break
                            if modified:
                                break
                        if modified:
                            break
                
                if not modified:
                    applied_changes.append(f"⚠️ SKIPPED REPLACE: Text not found in document: '{original_text_to_find[:30]}...'")
            
            elif action == "ADD" and new_text:
                # Add new content at the end of the document
                doc.add_paragraph(new_text)
                applied_changes.append(f"✅ ADDED to document: '{new_text[:50]}...'")
            
            elif action == "REMOVE" and original_text_to_find:
                modified = False
                # Search and remove from paragraphs
                for paragraph in doc.paragraphs:
                    if original_text_to_find in paragraph.text:
                        full_text = paragraph.text
                        new_full_text = full_text.replace(original_text_to_find, "")
                        paragraph.clear()
                        if new_full_text.strip():  # Only add back if there's remaining text
                            paragraph.add_run(new_full_text)
                        applied_changes.append(f"✅ REMOVED from document: '{original_text_to_find[:30]}...'")
                        modified = True
                        break
                
                if not modified:
                    applied_changes.append(f"⚠️ SKIPPED REMOVE: Text not found in document: '{original_text_to_find[:30]}...'")
        
        # Save modified document
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer, applied_changes
        
    except Exception as e:
        applied_changes.append(f"❌ ERROR modifying DOCX: {str(e)}")
        return None, applied_changes


def apply_modifications_to_pdf(
    original_file_bytes: bytes, 
    original_text: str, 
    modifications: List[Dict[str, str]], 
    filename: str
) -> Tuple[Optional[io.BytesIO], List[str]]:
    """Apply modifications to PDF document while preserving formatting"""
    applied_changes = []
    
    # For PDF, we need to create a new document since direct text editing is complex
    # We'll apply text modifications first, then create a styled PDF
    try:
        modified_text = original_text
        
        # Sort modifications by action priority
        action_priority = {"REMOVE": 1, "REPLACE": 2, "ADD": 3, "REORDER": 4}
        sorted_modifications = sorted(modifications, key=lambda x: action_priority.get(x.get("action", ""), 5))
        
        # Apply text modifications
        for mod in sorted_modifications:
            action = mod.get("action", "").upper()
            original_text_to_find = mod.get("original", "")
            new_text = mod.get("new", "")
            
            if action == "REPLACE" and original_text_to_find and new_text:
                if original_text_to_find in modified_text:
                    modified_text = modified_text.replace(original_text_to_find, new_text)
                    applied_changes.append(f"✅ REPLACED in PDF: '{original_text_to_find[:30]}...' with '{new_text[:30]}...'")
                else:
                    applied_changes.append(f"⚠️ SKIPPED REPLACE: Text not found: '{original_text_to_find[:30]}...'")
            
            elif action == "REMOVE" and original_text_to_find:
                if original_text_to_find in modified_text:
                    modified_text = modified_text.replace(original_text_to_find, "")
                    applied_changes.append(f"✅ REMOVED from PDF: '{original_text_to_find[:30]}...'")
                else:
                    applied_changes.append(f"⚠️ SKIPPED REMOVE: Text not found: '{original_text_to_find[:30]}...'")
            
            elif action == "ADD" and new_text:
                modified_text += "\n" + new_text
                applied_changes.append(f"✅ ADDED to PDF: '{new_text[:50]}...'")
        
        # Create new PDF with modified text using PyMuPDF (preserves some styling)
        if PYMUPDF_AVAILABLE:
            pdf_buffer = modify_pdf_with_pymupdf(original_file_bytes, modified_text, filename)
            if pdf_buffer:
                return pdf_buffer, applied_changes
        
        # Fallback to ReportLab
        if REPORTLAB_AVAILABLE:
            pdf_buffer = create_pdf_from_text_reportlab(modified_text, filename)
            if pdf_buffer:
                applied_changes.append("ℹ️ Created new PDF with ReportLab (basic styling)")
                return pdf_buffer, applied_changes
        
        applied_changes.append("❌ No PDF creation libraries available")
        return None, applied_changes
        
    except Exception as e:
        applied_changes.append(f"❌ ERROR modifying PDF: {str(e)}")
        return None, applied_changes


def apply_resume_modifications(original_resume: str, modifications: List[Dict[str, str]]) -> Tuple[str, List[str]]:
    """Apply the suggested modifications to the original resume text (fallback method)"""
    modified_resume = original_resume
    applied_changes = []
    
    # Sort modifications by action priority: REMOVE first, then REPLACE, then ADD, then REORDER
    action_priority = {"REMOVE": 1, "REPLACE": 2, "ADD": 3, "REORDER": 4}
    sorted_modifications = sorted(modifications, key=lambda x: action_priority.get(x.get("action", ""), 5))
    
    for mod in sorted_modifications:
        action = mod.get("action", "").upper()
        location = mod.get("location", "")
        original_text = mod.get("original", "")
        new_text = mod.get("new", "")
        reason = mod.get("reason", "")
        
        try:
            if action == "REPLACE" and original_text and new_text:
                if original_text in modified_resume:
                    modified_resume = modified_resume.replace(original_text, new_text)
                    applied_changes.append(f"✅ REPLACED: '{original_text[:50]}...' with '{new_text[:50]}...'")
                else:
                    applied_changes.append(f"⚠️ SKIPPED REPLACE: Original text not found: '{original_text[:50]}...'")
                    
            elif action == "REMOVE" and original_text:
                if original_text in modified_resume:
                    modified_resume = modified_resume.replace(original_text, "")
                    applied_changes.append(f"✅ REMOVED: '{original_text[:50]}...'")
                else:
                    applied_changes.append(f"⚠️ SKIPPED REMOVE: Text not found: '{original_text[:50]}...'")
                    
            elif action == "ADD" and new_text:
                # For ADD operations, we'll append to the end of the specified section or resume
                if location.lower() in ["end", "bottom", "resume end"]:
                    modified_resume += "\n" + new_text
                    applied_changes.append(f"✅ ADDED at end: '{new_text[:50]}...'")
                else:
                    # Try to find the section and add after it
                    section_patterns = [
                        f"{location}:",
                        location.upper() + ":",
                        location.title() + ":"
                    ]
                    
                    added = False
                    for pattern in section_patterns:
                        if pattern in modified_resume:
                            # Find the end of the section
                            section_start = modified_resume.find(pattern)
                            next_section = modified_resume.find("\n\n", section_start + len(pattern))
                            
                            if next_section == -1:
                                # Add at end of resume
                                modified_resume += "\n" + new_text
                            else:
                                # Insert before next section
                                modified_resume = modified_resume[:next_section] + "\n" + new_text + modified_resume[next_section:]
                            
                            applied_changes.append(f"✅ ADDED to {location}: '{new_text[:50]}...'")
                            added = True
                            break
                    
                    if not added:
                        # Fallback: add at end
                        modified_resume += "\n" + new_text
                        applied_changes.append(f"✅ ADDED at end (section not found): '{new_text[:50]}...'")
                        
            elif action == "REORDER":
                # REORDER is complex and depends on the specific implementation
                # For now, we'll note it but not implement automatic reordering
                applied_changes.append(f"ℹ️ REORDER suggested for {location}: {reason}")
                
            else:
                applied_changes.append(f"⚠️ UNSUPPORTED ACTION: {action}")
                
        except Exception as e:
            applied_changes.append(f"❌ ERROR applying {action}: {str(e)}")
    
    return modified_resume, applied_changes


def generate_customized_resume(
    original_resume: str,
    job_description: str,
    model: str = DEFAULT_MODEL,
) -> Dict[str, Any]:
    """Generate a customized resume by applying AI-suggested modifications to the original resume"""
    global request_in_progress, cancel_request

    try:
        # Set request in progress flag
        request_in_progress = True
        cancel_request = False

        # Check API server status
        api_server_running = check_api_server_status()

        # Create prompt
        prompt = create_resume_modification_prompt(original_resume, job_description)

        # Get client for API request
        client = get_ai_client()

        # Generate response
        result = ""
        if api_server_running and hasattr(client, "chat"):
            # Using OpenAI-compatible API
            st.info(f"Analyzing resume for improvements using AI model: {model}...")

            try:
                response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {
                            "role": "system", 
                            "content": "You are a professional resume writer and career coach. Analyze resumes and suggest specific, targeted modifications to improve job alignment while maintaining complete accuracy and truthfulness."
                        },
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.3,  # Lower temperature for more consistent, precise suggestions
                )

                result = response.choices[0].message.content
            except Exception as api_error:
                st.error(f"API error: {str(api_error)}")
                # Fall back to direct g4f
                api_server_running = False

        if not api_server_running:
            # Using g4f directly as fallback
            st.warning("Using g4f directly for resume analysis (this may take longer)...")

            try:
                import g4f

                g4f_model = get_g4f_model_from_name(model)
                if not g4f_model:
                    return {"success": False, "error": "Failed to get g4f model"}

                result = g4f.ChatCompletion.create(
                    model=g4f_model,
                    messages=[
                        {
                            "role": "system", 
                            "content": "You are a professional resume writer and career coach. Analyze resumes and suggest specific, targeted modifications to improve job alignment while maintaining complete accuracy and truthfulness."
                        },
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.3,
                )
            except Exception as g4f_error:
                st.error(f"Error using g4f directly: {str(g4f_error)}")
                return {"success": False, "error": str(g4f_error)}

        # Check if request was cancelled
        if cancel_request:
            return {"success": False, "error": "Request cancelled by user"}

        # Parse the modifications from AI response
        if result and result.strip():
            st.warning("Raw AI response:")
            st.text(result.strip())

            parse_result = parse_resume_modifications(result.strip())
            
            if not parse_result["success"]:
                return {"success": False, "error": f"Failed to parse modifications: {parse_result['error']}"}
            
            modifications_data = parse_result["data"]
            modifications = modifications_data.get("modifications", [])
            summary = modifications_data.get("summary", "")
            
            if not modifications:
                return {"success": False, "error": "No modifications suggested by AI"}
            
            # Apply modifications to the original resume text (for display purposes)
            st.info(f"Applying {len(modifications)} suggested modifications...")
            customized_resume, applied_changes = apply_resume_modifications(original_resume, modifications)
            
            # Show what changes were applied
            if applied_changes:
                st.info("Changes applied:")
                for change in applied_changes:
                    if change.startswith("✅"):
                        st.success(change)
                    elif change.startswith("⚠️"):
                        st.warning(change)
                    elif change.startswith("❌"):
                        st.error(change)
                    else:
                        st.info(change)
            
            return {
                "success": True, 
                "customized_resume": customized_resume,
                "modifications": modifications,
                "applied_changes": applied_changes,
                "summary": summary,
                "original_resume_text": original_resume,  # Store for document processing
                "message": f"Applied {len([c for c in applied_changes if c.startswith('✅')])} modifications successfully!"
            }
        else:
            return {"success": False, "error": "Empty response from AI model"}

    except Exception as e:
        st.error(f"Error in generate_customized_resume: {str(e)}")
        st.error(traceback.format_exc())
        return {"success": False, "error": str(e)}
    finally:
        # Reset request in progress flag
        request_in_progress = False


def create_complete_customized_resume_prompt(
    original_resume: str,
    job_description: str,
) -> str:
    """Create prompt for generating a complete customized resume with proper structure"""
    return f"""
    Create a complete, professionally structured resume that is specifically tailored to the job description provided. The goal is to optimize the resume for this specific role while maintaining complete truthfulness and accuracy.

    ## Original Resume:
    {original_resume}
    
    ## Job Description:
    {job_description}
    
    Please create a customized resume with the following structure:
    
    **Required Structure:**
    1. **Full Name** (from original resume)
    2. **Professional Job Title** - Create a targeted job title that aligns with the position being applied for, based on the candidate's actual experience and capabilities
    3. **Professional Summary** - Write a compelling 2-3 sentence summary that highlights the candidate's most relevant experience, skills, and achievements for this specific role. Use keywords from the job description where appropriate.
    4. **Contact Information** (if present in original)
    5. **Core Skills/Technical Skills** - Reorganize and highlight the most relevant skills for this position
    6. **Professional Experience** - Reorder and emphasize the most relevant experience, optimizing bullet points with job-specific keywords
    7. **Education** - Include education information as in original
    8. **Additional sections** - Include any other sections from original resume (certifications, projects, etc.)
    
    **Key Requirements:**
    - **Job Title**: Must reflect what the candidate can actually do based on their experience (e.g., if they have 3+ years of software development experience and are applying for a "Senior Software Engineer" role, use that title)
    - **Professional Summary**: Must be specific to this role, compelling, and include relevant keywords from the job description
    - **Content Truthfulness**: Do NOT add any experience, skills, or achievements that weren't in the original resume
    - **Keyword Optimization**: Use industry-standard keywords from the job description where they accurately reflect the candidate's experience
    - **Relevance Prioritization**: Reorder sections and bullet points to emphasize the most relevant information first
    - **Professional Presentation**: Maintain professional tone and clear structure
    
    **Examples of Good Job Titles based on experience:**
    - "Software Engineer" → "Full Stack Developer" (if applying for full stack role)
    - "Marketing Coordinator" → "Digital Marketing Specialist" (if applying for digital marketing role)
    - "Data Analyst" → "Senior Data Analyst" (if they have 3+ years experience)
    
    **Professional Summary Examples:**
    - "Results-driven Software Engineer with 4+ years of experience developing scalable web applications using React, Node.js, and Python. Proven track record of delivering high-quality solutions in Agile environments and collaborating effectively with cross-functional teams."
    - "Digital Marketing Professional with expertise in SEO, content marketing, and social media strategy. Successfully increased organic traffic by 150% and generated $2M in revenue through integrated marketing campaigns."
    
    Return ONLY the complete customized resume text, properly formatted and ready to be saved as a document. Do not include any additional explanations or comments.
    """


def generate_complete_customized_resume(
    original_resume: str,
    job_description: str,
    model: str = DEFAULT_MODEL,
) -> Dict[str, Any]:
    """Generate a complete customized resume with proper structure including job title and summary"""
    global request_in_progress, cancel_request

    try:
        # Set request in progress flag
        request_in_progress = True
        cancel_request = False

        # Check API server status
        api_server_running = check_api_server_status()

        # Create prompt for complete resume generation
        prompt = create_complete_customized_resume_prompt(original_resume, job_description)

        # Get client for API request
        client = get_ai_client()

        # Generate response
        result = ""
        if api_server_running and hasattr(client, "chat"):
            # Using OpenAI-compatible API
            st.info(f"Generating complete customized resume using AI model: {model}...")

            try:
                response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {
                            "role": "system", 
                            "content": "You are a professional resume writer and career coach. Create complete, well-structured resumes that are specifically tailored to job descriptions while maintaining complete accuracy and truthfulness. Always include a targeted job title and compelling professional summary."
                        },
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.3,  # Lower temperature for more consistent, precise results
                )

                result = response.choices[0].message.content
            except Exception as api_error:
                st.error(f"API error: {str(api_error)}")
                # Fall back to direct g4f
                api_server_running = False

        if not api_server_running:
            # Using g4f directly as fallback
            st.warning("Using g4f directly for resume generation (this may take longer)...")

            try:
                import g4f

                g4f_model = get_g4f_model_from_name(model)
                if not g4f_model:
                    return {"success": False, "error": "Failed to get g4f model"}

                result = g4f.ChatCompletion.create(
                    model=g4f_model,
                    messages=[
                        {
                            "role": "system", 
                            "content": "You are a professional resume writer and career coach. Create complete, well-structured resumes that are specifically tailored to job descriptions while maintaining complete accuracy and truthfulness. Always include a targeted job title and compelling professional summary."
                        },
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.3,
                )
            except Exception as g4f_error:
                st.error(f"Error using g4f directly: {str(g4f_error)}")
                return {"success": False, "error": str(g4f_error)}

        # Check if request was cancelled
        if cancel_request:
            return {"success": False, "error": "Request cancelled by user"}

        # Return the complete customized resume
        if result and result.strip():
            customized_resume = result.strip()
            
            return {
                "success": True, 
                "customized_resume": customized_resume,
                "message": "Complete customized resume generated successfully with targeted job title and professional summary!"
            }
        else:
            return {"success": False, "error": "Empty response from AI model"}

    except Exception as e:
        st.error(f"Error in generate_complete_customized_resume: {str(e)}")
        st.error(traceback.format_exc())
        return {"success": False, "error": str(e)}
    finally:
        # Reset request in progress flag
        request_in_progress = False
