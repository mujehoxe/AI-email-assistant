import streamlit as st
from utils import extract_name, send_email
from template_manager import save_template
from ai_helper import (
    generate_improved_template,
    generate_customized_resume,
    modify_original_resume_document,
    get_resume_file_extension,
    start_g4f_api_server,
    cancel_ai_request,
    is_request_in_progress,
    AVAILABLE_MODELS,
    get_available_models,
    check_api_server_status,
    extract_contact_info,
    DEFAULT_MODEL,
    extract_company_name,
)
import io
import traceback
import os
import re
import json
import pickle
import tempfile
import base64
from pathlib import Path

# Try importing PDF and docx libraries
try:
    import PyPDF2
    import docx

    resume_parsing_available = True
except ImportError:
    st.warning(
        "Resume parsing libraries not available. Install PyPDF2 and python-docx for resume parsing."
    )
    resume_parsing_available = False

# Create data directory if it doesn't exist
os.makedirs("data", exist_ok=True)
RESUME_STORAGE_PATH = "data/resume_data.pkl"


def save_resume_to_disk(filename, content, file_bytes=None):
    """Save resume data to disk for persistence"""
    data = {"filename": filename, "content": content, "file_bytes": file_bytes}
    with open(RESUME_STORAGE_PATH, "wb") as f:
        pickle.dump(data, f)


def load_resume_from_disk():
    """Load resume data from disk"""
    if not os.path.exists(RESUME_STORAGE_PATH):
        return None

    try:
        with open(RESUME_STORAGE_PATH, "rb") as f:
            return pickle.load(f)
    except:
        return None


def extract_text_from_resume(resume_file):
    """Extract text content from resume file"""
    if resume_file is None or not resume_parsing_available:
        return ""

    try:
        file_extension = resume_file.name.split(".")[-1].lower()

        if file_extension == "pdf":
            # Extract text from PDF
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(resume_file.getvalue()))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text

        elif file_extension in ["docx", "doc"]:
            # Extract text from Word document
            doc = docx.Document(io.BytesIO(resume_file.getvalue()))
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text

        else:
            # For other formats, return empty string
            st.warning(
                f"Unsupported file format: {file_extension}. Only PDF and DOCX are supported."
            )
            return ""

    except Exception as e:
        st.error(f"Error extracting text from resume: {str(e)}")
        st.error(traceback.format_exc())
        return ""


def infer_smtp_server(email):
    """Infer SMTP server from email address"""
    if not email:
        return ""

    domain = email.split("@")[-1].lower()

    # Common email providers
    smtp_servers = {
        "gmail.com": "smtp.gmail.com",
        "outlook.com": "smtp.office365.com",
        "hotmail.com": "smtp.office365.com",
        "yahoo.com": "smtp.mail.yahoo.com",
        "aol.com": "smtp.aol.com",
        "icloud.com": "smtp.mail.me.com",
        "protonmail.com": "smtp.protonmail.ch",
    }

    return smtp_servers.get(domain, f"mail.{domain}")


def render_email_settings_sidebar():
    """Render email settings in sidebar"""
    st.sidebar.header("Email Settings")

    # Get environment variables if they exist
    default_email = os.environ.get("EMAIL_ADDRESS", "")
    default_password = os.environ.get("EMAIL_PASSWORD", "")

    # Email credentials
    email = st.sidebar.text_input("Your Email Address", value=default_email)
    if email:
        # Store in session state
        st.session_state.sender_email = email

    password = st.sidebar.text_input(
        "Email Password", type="password", value=default_password
    )
    if password:
        # Store in session state
        st.session_state.sender_password = password

    # SMTP server settings
    inferred_server = infer_smtp_server(email) if email else ""
    smtp_server = st.sidebar.text_input("SMTP Server", value=inferred_server)
    if smtp_server:
        st.session_state.smtp_server = smtp_server

    smtp_port = st.sidebar.selectbox(
        "SMTP Port",
        options=[587, 465, 25],
        index=0,
        help="Port 587 uses STARTTLS, 465 uses SSL/TLS",
    )
    st.session_state.smtp_port = smtp_port

    # Save settings to .env file
    if st.sidebar.button("Save Email Settings"):
        try:
            with open(".env", "w") as f:
                f.write(f"EMAIL_ADDRESS={email}\n")
                f.write(f"EMAIL_PASSWORD={password}\n")
                f.write(f"SMTP_SERVER={smtp_server}\n")
                f.write(f"SMTP_PORT={smtp_port}\n")
            st.sidebar.success("Email settings saved to .env file")
        except Exception as e:
            st.sidebar.error(f"Error saving settings: {str(e)}")

    # Add G4F API server controls to sidebar
    st.sidebar.header("AI Server Settings")

    # Check if API server is running
    server_status = check_api_server_status()
    status_color = "🟢" if server_status else "🔴"
    st.sidebar.write(
        f"{status_color} Server status: {'Running' if server_status else 'Stopped'}"
    )

    if st.sidebar.button("Start G4F API Server"):
        try:
            start_g4f_api_server()
            st.sidebar.success("G4F API server is running")
            # Force fetch the models after starting the server
            try:
                models = get_available_models()
                if models and models != AVAILABLE_MODELS:
                    st.session_state.ai_models = models
                    st.sidebar.success(f"Found {len(models)} models")
                    st.rerun()
            except Exception as e:
                st.sidebar.error(f"Error fetching models: {str(e)}")
        except Exception as e:
            st.sidebar.error(f"Error starting G4F API server: {str(e)}")
            st.sidebar.error(traceback.format_exc())

    if st.sidebar.button("Refresh AI Models"):
        try:
            models = get_available_models()
            st.session_state.ai_models = models
            st.sidebar.success(f"Found {len(models)} models")
            st.rerun()
        except Exception as e:
            st.sidebar.error(f"Error refreshing models: {str(e)}")


def render_email_form():
    """Render the email form with template generation and sending functionality"""
    from ai_helper import (
        get_available_models,
        generate_improved_template,
        check_api_server_status,
        is_request_in_progress,
        cancel_ai_request,
        extract_contact_info,
        DEFAULT_MODEL,
    )
    from template_manager import save_template
    from utils import send_email, infer_smtp_settings

    # Initialize session state variables if they don't exist
    if "recipient_email" not in st.session_state:
        st.session_state.recipient_email = ""
    if "recipient_name" not in st.session_state:
        st.session_state.recipient_name = ""
    if "company_name" not in st.session_state:
        st.session_state.company_name = ""
    if "position" not in st.session_state:
        st.session_state.position = ""
    if "subject" not in st.session_state:
        st.session_state.subject = "Job Application"
    if "resume_content" not in st.session_state:
        st.session_state.resume_content = ""
    if "resume_filename" not in st.session_state:
        st.session_state.resume_filename = ""
    if "customized_resume" not in st.session_state:
        st.session_state.customized_resume = ""
    if "use_customized_resume" not in st.session_state:
        st.session_state.use_customized_resume = False
    if "customized_resume_docx" not in st.session_state:
        st.session_state.customized_resume_docx = None
    if "customized_resume_filename" not in st.session_state:
        st.session_state.customized_resume_filename = ""

    # Render email settings in sidebar
    render_email_settings_sidebar()

    # Create multi-page layout using tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📄 Resume Upload", 
        "📝 Job Description", 
        "🎯 Resume Customization", 
        "📧 Template Generation", 
        "✉️ Email Sending"
    ])

    with tab1:
        render_resume_upload_tab()
    
    with tab2:
        render_job_description_tab()
    
    with tab3:
        render_resume_customization_tab()
    
    with tab4:
        render_template_generation_tab()
    
    with tab5:
        render_email_sending_tab()


def render_resume_upload_tab():
    """Render the resume upload and processing tab"""
    st.header("Resume Upload & Processing")
    
    # Try to load resume data from disk at the start
    if not st.session_state.resume_content:
        resume_data = load_resume_from_disk()
        if resume_data:
            st.session_state.resume_filename = resume_data["filename"]
            st.session_state.resume_content = resume_data["content"]
            st.session_state.resume_file_bytes = resume_data.get("file_bytes", None)

    # Display currently loaded resume if any
    if st.session_state.resume_filename:
        st.info(f"Currently loaded resume: {st.session_state.resume_filename}")

        # Add option to clear the loaded resume
        if st.button("Clear Resume", key="clear_resume_upload"):
            st.session_state.resume_content = ""
            st.session_state.resume_filename = ""
            st.session_state.resume_file_bytes = None
            st.session_state.customized_resume = ""
            st.session_state.customized_resume_docx = None
            st.session_state.customized_resume_filename = ""
            st.session_state.use_customized_resume = False

            # Remove the disk file too
            if os.path.exists(RESUME_STORAGE_PATH):
                os.remove(RESUME_STORAGE_PATH)

            st.rerun()

    # File uploader with key to ensure proper state tracking
    resume_file = st.file_uploader(
        "Upload your resume (PDF or DOCX)",
        type=["pdf", "docx"],
        key="resume_uploader",
    )

    # Process the file if a new one was uploaded
    if resume_file is not None:
        # Check if this is a new file upload or different from stored file
        if (
            not st.session_state.resume_filename
            or st.session_state.resume_filename != resume_file.name
        ):
            try:
                if resume_file.name.endswith(".pdf"):
                    from PyPDF2 import PdfReader

                    pdf = PdfReader(resume_file)
                    resume_content = ""
                    for page in pdf.pages:
                        resume_content += page.extract_text()
                elif resume_file.name.endswith(".docx"):
                    from docx import Document

                    doc = Document(resume_file)
                    resume_content = "\n".join(
                        [para.text for para in doc.paragraphs]
                    )

                # Store file content and name in session state
                st.session_state.resume_content = resume_content
                st.session_state.resume_filename = resume_file.name

                # Store the raw file bytes for sending as attachment
                file_bytes = resume_file.getvalue()
                st.session_state.resume_file_bytes = file_bytes

                # Save to disk for persistence
                save_resume_to_disk(resume_file.name, resume_content, file_bytes)

                st.success(f"Resume processed: {resume_file.name}")
            except Exception as e:
                st.error(f"Error processing resume: {str(e)}")

    # Show resume preview if we have content (from any source)
    if st.session_state.resume_content and resume_parsing_available:
        with st.expander("Resume Preview"):
            preview = (
                st.session_state.resume_content[:1000] + "..."
                if len(st.session_state.resume_content) > 1000
                else st.session_state.resume_content
            )
            st.text_area("Extracted Text", value=preview, height=150)

    # Status and next steps
    if st.session_state.resume_content:
        st.success("✅ Resume loaded successfully! You can now proceed to the next tab.")
    else:
        st.info("📤 Please upload your resume to continue.")


def render_job_description_tab():
    """Render the job description input tab"""
    st.header("Job Description")
    
    # Check if resume is uploaded
    if not st.session_state.resume_content:
        st.warning("⚠️ Please upload your resume in the 'Resume Upload' tab first.")
        return

    st.info("📝 Enter the job description for the position you're applying to. This will be used to customize your resume and generate a personalized email template.")
    
    job_description = st.text_area(
        "Paste job description here",
        value=st.session_state.get("job_description", ""),
        height=400,
        key="job_description",
        help="Copy and paste the complete job description, including requirements, responsibilities, and company information."
    )
    
    # Status and next steps
    if job_description:
        st.success("✅ Job description added! You can now proceed to resume customization.")
        
        # Show some extracted information preview
        if len(job_description) > 100:
            st.markdown("#### Quick Analysis")
            
            # Extract some basic info for preview
            from ai_helper import extract_contact_info
            contact_info = extract_contact_info(job_description)
            
            col1, col2, col3 = st.columns(3)
            with col1:
                if contact_info.get('company'):
                    st.metric("Company Detected", contact_info['company'])
                else:
                    st.metric("Company Detected", "Not found")
            
            with col2:
                if contact_info.get('email'):
                    st.metric("Contact Email", "Found")
                else:
                    st.metric("Contact Email", "Not found")
            
            with col3:
                word_count = len(job_description.split())
                st.metric("Word Count", f"{word_count} words")
    else:
        st.info("📝 Please paste the job description to continue.")


def render_resume_customization_tab():
    """Render the resume customization tab"""
    st.header("Resume Customization")
    
    # Check prerequisites
    if not st.session_state.resume_content:
        st.warning("⚠️ Please upload your resume in the 'Resume Upload' tab first.")
        return
    
    job_description = st.session_state.get("job_description", "")
    if not job_description:
        st.warning("⚠️ Please add a job description in the 'Job Description' tab first.")
        return

    st.info("🎯 Create a customized version of your resume that's tailored specifically to this job description. This will highlight your most relevant experience and skills.")
    
    # Resume customization controls
    customize_col, customize_cancel_col = st.columns([3, 1])
    
    with customize_col:
        if st.button(
            "Generate Customized Resume",
            disabled=is_request_in_progress(),
            help="Create a version of your resume tailored to this job description",
            type="primary",
            key="generate_customized_resume"
        ):
            with st.spinner("Customizing your resume..."):
                # Get the selected model (same as email generation)
                if "ai_models" not in st.session_state:
                    st.session_state.ai_models = get_available_models()
                
                # Find default model index
                default_index = 0
                if DEFAULT_MODEL in st.session_state.ai_models:
                    default_index = st.session_state.ai_models.index(DEFAULT_MODEL)
                
                selected_model = st.session_state.ai_models[default_index] if st.session_state.ai_models else DEFAULT_MODEL
                
                # Generate complete customized resume with proper structure
                from ai_helper import generate_complete_customized_resume
                
                result = generate_complete_customized_resume(
                    st.session_state.resume_content,
                    job_description,
                    model=selected_model
                )
                
                if result.get("success", False):
                    st.session_state.customized_resume = result["customized_resume"]
                    st.session_state.use_customized_resume = True
                    
                    # Now create the customized document using the original file
                    if hasattr(st.session_state, 'resume_file_bytes') and st.session_state.resume_file_bytes:
                        from ai_helper import modify_original_resume_document
                        
                        # Modify the original document while preserving styling
                        customized_doc = modify_original_resume_document(
                            st.session_state.resume_file_bytes,
                            st.session_state.customized_resume,
                            st.session_state.resume_filename
                        )
                        
                        if customized_doc:
                            st.session_state.customized_resume_docx = customized_doc
                            base_name = st.session_state.resume_filename.rsplit('.', 1)[0]
                            original_ext = get_resume_file_extension(st.session_state.resume_filename)
                            st.session_state.customized_resume_filename = f"{base_name}_customized.{original_ext}"
                            st.success("📄 Professional document created preserving original styling!")
                        else:
                            st.warning("⚠️ Could not create styled document, but customized text is available.")
                    else:
                        st.warning("⚠️ Original file not available for styling preservation. Text customization completed.")
                    
                    st.success("✅ Customized resume generated successfully!")
                    st.rerun()
                else:
                    st.error(f"❌ Failed to generate customized resume: {result.get('error', 'Unknown error')}")
    
    with customize_cancel_col:
        if st.button("Cancel", disabled=not is_request_in_progress(), key="cancel_resume_customization"):
            if cancel_ai_request():
                st.warning("Resume generation cancelled")
                st.rerun()

    # Show customized resume if available
    if st.session_state.customized_resume:
        st.markdown("### ✅ Customized Resume Generated")
        
        # Toggle between original and customized resume
        resume_choice = st.radio(
            "Choose which resume version to use:",
            options=["Use Original Resume", "Use Customized Resume"],
            index=1 if st.session_state.use_customized_resume else 0,
            horizontal=True
        )
        
        st.session_state.use_customized_resume = (resume_choice == "Use Customized Resume")
        
        # Prominent Download Section
        st.markdown("### 📥 Download Your Customized Resume")
        
        # Download buttons in prominent location
        col1, col2, col3 = st.columns([1, 1, 1])
        
        with col1:
            st.download_button(
                label="📄 Download as Text",
                data=st.session_state.customized_resume,
                file_name="customized_resume.txt",
                mime="text/plain",
                help="Download the customized resume as a text file",
                use_container_width=True
            )
        
        with col2:
            # Show styled document download button if available
            if st.session_state.customized_resume_docx and st.session_state.customized_resume_filename:
                # Determine file type and appropriate mime type
                file_ext = get_resume_file_extension(st.session_state.customized_resume_filename)
                if file_ext == "pdf":
                    mime_type = "application/pdf"
                    label = "📄 Download as PDF"
                    help_text = "Download the customized resume as a styled PDF document"
                else:  # docx
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    label = "📄 Download as Word"
                    help_text = "Download the customized resume as a styled Word document"
                
                st.download_button(
                    label=label,
                    data=st.session_state.customized_resume_docx.getvalue(),
                    file_name=st.session_state.customized_resume_filename,
                    mime=mime_type,
                    help=help_text,
                    use_container_width=True
                )
            else:
                st.info("📄 Styled document not available")
        
        with col3:
            if st.button("🔄 Regenerate Resume", help="Generate a new customized version", use_container_width=True):
                st.session_state.customized_resume = ""
                st.session_state.customized_resume_docx = None
                st.session_state.customized_resume_filename = ""
                st.session_state.use_customized_resume = False
                st.rerun()
        
        
        # Text Preview Section (Collapsible)
        with st.expander("📝 View Text Content", expanded=False):
            st.text_area(
                "Customized Resume Text Content",
                value=st.session_state.customized_resume,
                height=400,
                disabled=True,
                help="This is the text content used to generate your styled document"
            )
        
        # Option to clear customized resume
        if st.button("🗑️ Clear Customized Resume", key="clear_customized_resume"):
            st.session_state.customized_resume = ""
            st.session_state.customized_resume_docx = None
            st.session_state.customized_resume_filename = ""
            st.session_state.use_customized_resume = False
            st.rerun()
        
        # Status
        st.success("✅ Customized resume ready! You can now proceed to template generation.")
    
    else:
        st.info("👆 Click 'Generate Customized Resume' to create a tailored version of your resume for this job.")


def render_template_generation_tab():
    """Render the email template generation tab"""
    st.header("Email Template Generation")
    
    # Check prerequisites
    if not st.session_state.resume_content:
        st.warning("⚠️ Please upload your resume in the 'Resume Upload' tab first.")
        return
    
    job_description = st.session_state.get("job_description", "")
    if not job_description:
        st.warning("⚠️ Please add a job description in the 'Job Description' tab first.")
        return

    st.info("📧 Generate a personalized email template using AI. Choose whether to use your original or customized resume.")
    
    # Show which resume will be used
    if st.session_state.customized_resume:
        if st.session_state.use_customized_resume:
            st.success("🎯 Using customized resume for template generation")
        else:
            st.info("📄 Using original resume for template generation")
        
        # Allow user to change resume choice
        resume_choice = st.radio(
            "Which resume version do you want to use?",
            options=["Use Original Resume", "Use Customized Resume"],
            index=1 if st.session_state.use_customized_resume else 0,
            horizontal=True
        )
        st.session_state.use_customized_resume = (resume_choice == "Use Customized Resume")
    else:
        st.info("📄 Using original resume (no customized version available)")

    # AI Model Configuration
    st.markdown("### AI Model Configuration")
    
    # Check API server status
    server_status = "Not Running"
    if check_api_server_status():
        server_status = "Running"

    col1, col2 = st.columns([2, 1])
    with col1:
        st.markdown(f"**API Server Status:** {server_status}")

    # Get available models
    if "ai_models" not in st.session_state:
        st.session_state.ai_models = get_available_models()

    model_count = len(st.session_state.ai_models)
    st.markdown(f"**Available Models:** {model_count}")

    # Find default model index
    default_index = 0
    if DEFAULT_MODEL in st.session_state.ai_models:
        default_index = st.session_state.ai_models.index(DEFAULT_MODEL)

    # Model selection
    selected_model = st.selectbox(
        "Select AI Model",
        options=st.session_state.ai_models,
        index=default_index,
    )

    # Refresh models button
    if st.button("Refresh Models List", key="refresh_models_list"):
        st.session_state.ai_models = get_available_models()
        st.rerun()

    # Generate template section
    st.markdown("### Generate Email Template")
    
    # Generate template button
    generate_col, cancel_col = st.columns([3, 1])

    with generate_col:
        if generate_button := st.button(
            "Generate Email Template",
            disabled=is_request_in_progress(),
            type="primary",
            key="generate_email_template"
        ):
            with st.spinner("Generating personalized email template..."):
                # Get current template for reference
                current_template = st.session_state.current_template

                # Choose which resume to use based on user preference
                resume_to_use = st.session_state.resume_content
                if st.session_state.use_customized_resume and st.session_state.customized_resume:
                    resume_to_use = st.session_state.customized_resume

                # Generate template using selected resume content
                result = generate_improved_template(
                    job_description,
                    resume_to_use,
                    current_template,
                    model=selected_model,
                )

                # Store result in session state
                st.session_state.result = result

                if result.get("success", False):
                    template = result["template"]

                    # Update template in session state
                    st.session_state.current_template = {
                        "greeting": template.get("greeting", ""),
                        "body": template.get("body", ""),
                        "signature": template.get("signature", ""),
                    }

                    # Update other extracted information if they exist and are not empty
                    if template.get("position"):
                        st.session_state.position = template["position"]
                        st.success(f"✅ Position extracted: {template['position']}")

                    if template.get("employer") and template["employer"]:
                        st.session_state.company_name = template["employer"]
                        st.success(f"✅ Company name extracted: {template['employer']}")
                    elif template.get("company") and template["company"]:
                        st.session_state.company_name = template["company"]
                        st.success(f"✅ Company name extracted: {template['company']}")

                    if template.get("subject"):
                        st.session_state.subject = template["subject"]

                    if template.get("recipient_email"):
                        st.session_state.recipient_email = template["recipient_email"]
                        st.info(f"📧 Contact email extracted: {template['recipient_email']}")

                    if template.get("recipient_name"):
                        st.session_state.recipient_name = template.get("recipient_name")
                        st.info(f"👤 Contact name extracted: {template.get('recipient_name')}")

                    st.success("🎉 Email template generated successfully! You can now proceed to email sending.")
                    st.rerun()
                else:
                    st.error(f"❌ Failed to generate template: {result.get('error', 'Unknown error')}")

    with cancel_col:
        if st.button("Cancel", disabled=not is_request_in_progress(), key="cancel_template_generation"):
            if cancel_ai_request():
                st.warning("Request cancelled")
                st.rerun()

    # Show template preview if available
    if st.session_state.current_template.get("body"):
        st.markdown("### Generated Template Preview")
        
        with st.expander("📧 View Generated Email Template", expanded=True):
            st.markdown("**Greeting:**")
            st.text(st.session_state.current_template.get("greeting", ""))
            
            st.markdown("**Body:**")
            st.text(st.session_state.current_template.get("body", ""))
            
            st.markdown("**Signature:**")
            st.text(st.session_state.current_template.get("signature", ""))
        
        st.success("✅ Template ready! Proceed to the 'Email Sending' tab to customize and send your email.")


def render_email_sending_tab():
    """Render the email sending and configuration tab"""
    st.header("Email Sending & Configuration")
    
    # Check prerequisites
    if not st.session_state.resume_content:
        st.warning("⚠️ Please upload your resume in the 'Resume Upload' tab first.")
        return
    
    # Check if we have a template loaded - if yes, skip job description requirement
    if not st.session_state.current_template.get("body"):
        st.warning("⚠️ Please generate an email template in the 'Template Generation' tab first, or load a template from the sidebar.")
        return
    
    # Only require job description if no template is loaded
    job_description = st.session_state.get("job_description", "")
    if not job_description and not st.session_state.current_template.get("body"):
        st.warning("⚠️ Please add a job description in the 'Job Description' tab first.")
        return

    # Create layout
    col1, col2 = st.columns([3, 2])

    with col1:
        # Email Configuration section
        st.markdown("### Email Configuration")

        # Recipient information
        recipient_email = st.text_input(
            "Recipient Email",
            value=st.session_state.recipient_email,
            help="Enter the recruiter's email address",
        )
        st.session_state.recipient_email = recipient_email

        recipient_name = st.text_input(
            "Recipient Name",
            value=st.session_state.recipient_name,
            help="Enter the recipient's name (if left empty, will try to extract from email)",
        )
        st.session_state.recipient_name = recipient_name

        company_name = st.text_input(
            "Company Name",
            value=st.session_state.company_name,
            help="Enter the company name you're applying to",
        )
        st.session_state.company_name = company_name

        position = st.text_input(
            "Position/Designation",
            value=st.session_state.position,
            help="Enter the position you're applying for",
        )
        st.session_state.position = position

    with col2:
        # Email preview with placeholders filled in
        st.header("Email Preview")

        # Get values for placeholders
        recipient_name = st.session_state.recipient_name
        company_name = st.session_state.company_name
        position_name = st.session_state.position

        # Don't extract from email - only use job description data
        # If recipient_name is empty, we'll use the fallback

        # Extract first name for more personalized greeting
        first_name = ""
        if recipient_name:
            name_parts = recipient_name.split()
            if name_parts:
                first_name = name_parts[0]  # Get first name

        # Use "Hiring Manager" if no valid name is found
        display_name = first_name if first_name else "Hiring Manager"

        # Update subject with position if available
        if (
            position_name
            and company_name
            and not st.session_state.get("subject_manually_set", False)
        ):
            subject_value = (
                f"Application for the {position_name} position at {company_name}"
            )
        elif position_name and not st.session_state.get("subject_manually_set", False):
            subject_value = f"Application for the {position_name} position"
        else:
            subject_value = st.session_state.subject

        # MOVED: Subject field to top of preview section
        st.markdown("### Subject")
        subject = st.text_input(
            "Email Subject", value=subject_value, key="email_subject"
        )

        # Track if subject was manually set
        if subject != subject_value:
            st.session_state.subject_manually_set = True
        st.session_state.subject = subject

        # Replace placeholders in template
        greeting = st.session_state.current_template["greeting"]
        if "{name}" in greeting:
            # Use first name if available, otherwise use full name or fallback
            greeting = greeting.replace("{name}", display_name)

        body = st.session_state.current_template["body"]
        if "{position}" in body:
            body = body.replace("{position}", position_name or "position")
        if "{company}" in body:
            # Only replace {company} with actual company name if we have one, otherwise leave it blank
            if company_name:
                body = body.replace("{company}", company_name)
            else:
                body = body.replace("{company}", "")

        signature = st.session_state.current_template["signature"]

        # Display editable preview
        st.markdown("### Greeting")
        preview_greeting = st.text_area(
            "Edit greeting",
            greeting,
            height=68,
            key="preview_greeting",
        )

        st.markdown("### Body")
        preview_body = st.text_area(
            "Edit body",
            body,
            height=200,
            key="preview_body",
        )

        st.markdown("### Signature")
        preview_signature = st.text_area(
            "Edit signature",
            signature,
            height=100,
            key="preview_signature",
        )

        # Add action buttons
        col1, col2 = st.columns(2)

        # Add Send Email button
        with col1:
            if st.button("Send Email", type="primary", key="send_email"):
                if not recipient_email:
                    st.error("Please enter a recipient email address.")
                elif not st.session_state.get("sender_email"):
                    st.error("Please configure your email address in the sidebar.")
                elif not st.session_state.get("sender_password"):
                    st.error("Please configure your email password in the sidebar.")
                else:
                    # Get the resume file or create a temp file if we have cached content
                    attachment = None

                    # If user chose to use customized resume and it exists
                    if st.session_state.use_customized_resume and st.session_state.customized_resume:
                        # Check if we have a Word document version
                        if st.session_state.customized_resume_docx and st.session_state.customized_resume_filename:
                            attachment = st.session_state.customized_resume_docx
                            attachment.name = st.session_state.customized_resume_filename
                            st.info("📧 Attaching customized resume (Word document) to email")
                        else:
                            # Fall back to text version
                            from io import BytesIO
                            attachment = BytesIO(st.session_state.customized_resume.encode("utf-8"))
                            base_name = st.session_state.resume_filename.rsplit('.', 1)[0] if st.session_state.resume_filename else "resume"
                            attachment.name = f"{base_name}_customized.txt"
                            st.info("📧 Attaching customized resume (text) to email")
                    
                    # Otherwise, use original resume
                    elif (
                        st.session_state.get("resume_file_bytes")
                        and st.session_state.resume_filename
                    ):
                        from io import BytesIO
                        attachment = BytesIO(st.session_state.resume_file_bytes)
                        attachment.name = st.session_state.resume_filename
                        st.info("📧 Attaching original resume to email")

                    # As a last resort, if we only have the extracted text but no binary data
                    elif (
                        st.session_state.resume_content
                        and st.session_state.resume_filename
                    ):
                        from io import BytesIO
                        attachment = BytesIO(st.session_state.resume_content.encode("utf-8"))
                        attachment.name = st.session_state.resume_filename
                        st.info("📧 Attaching original resume to email")

                    success, message = send_email(
                        recipient_email=recipient_email,
                        subject=subject,
                        greeting=preview_greeting,
                        body=preview_body,
                        signature=preview_signature,
                        attachment=attachment,
                        sender_email=st.session_state.get("sender_email"),
                        sender_password=st.session_state.get("sender_password"),
                        smtp_server=st.session_state.get("smtp_server"),
                        smtp_port=st.session_state.get("smtp_port", 587),
                    )

                    if success:
                        st.success(message)
                    else:
                        st.error(message)

        # Copy to clipboard button
        with col2:
            if st.button("Copy Template to Clipboard", key="copy_template"):
                full_email = (
                    f"{preview_greeting}\n\n{preview_body}\n\n{preview_signature}"
                )

                # Prepare the email content for JavaScript
                # Use json.dumps to properly escape all special characters
                escaped_email = json.dumps(full_email)

                # Use JavaScript to copy to clipboard
                js_code = f"""
                <script>
                    const text = {escaped_email};
                    const el = document.createElement('textarea');
                    el.value = text;
                    document.body.appendChild(el);
                    el.select();
                    document.execCommand('copy');
                    document.body.removeChild(el);
                    alert('Email template copied to clipboard!');
                </script>
                """
                st.components.v1.html(js_code, height=0)
                st.success("Email template copied to clipboard!")

        # Template save form
        st.markdown("### Save Template")
        template_name = st.text_input("Template Name", "")

        if st.button("Save Template", disabled=not template_name, key="save_template"):
            if save_template(
                template_name, preview_greeting, preview_body, preview_signature
            ):
                st.success(f"Template '{template_name}' saved!")
            else:
                st.error("Failed to save template")
